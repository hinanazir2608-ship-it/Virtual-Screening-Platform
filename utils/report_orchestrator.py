import os
import json
import docx
from docx.shared import Inches, Pt, RGBColor
from typing import Dict, Any, List
from utils.results_engine import scan_and_rank_results, write_results_json, write_results_csv
from utils.rag_handler import LiteratureRAGHandler

try:
    from rdkit import Chem
    from rdkit.Chem import Descriptors, Lipinski
    RDKIT_AVAILABLE = True
except ImportError:
    RDKIT_AVAILABLE = False

def run_admet_analysis(mol_path: str) -> Dict[str, Any]:
    """Computes basic Lipinski Rule of 5 ADMET descriptors using RDKit."""
    if not RDKIT_AVAILABLE or not os.path.exists(mol_path):
        return {"MW": "N/A", "LogP": "N/A", "HBD": "N/A", "HBA": "N/A", "Violations": "N/A"}

    try:
        # PyTorch update: Read SDF or PDB instead of PDBQT
        if mol_path.endswith('.sdf'):
            mol = Chem.MolFromMolFile(mol_path)
        elif mol_path.endswith('.pdb'):
            mol = Chem.MolFromPDBFile(mol_path)
        else:
            mol = Chem.MolFromSmiles(mol_path)

        if not mol:
            return {"MW": "N/A", "LogP": "N/A", "HBD": "N/A", "HBA": "N/A", "Violations": "N/A"}

        mw = Descriptors.MolWt(mol)
        logp = Descriptors.MolLogP(mol)
        hbd = Lipinski.NumHDonors(mol)
        hba = Lipinski.NumHAcceptors(mol)
        
        violations = sum([mw > 500, logp > 5, hbd > 5, hba > 10])

        return {
            "MW": round(mw, 2),
            "LogP": round(logp, 2),
            "HBD": hbd,
            "HBA": hba,
            "Violations": violations
        }
    except Exception:
        return {"MW": "N/A", "LogP": "N/A", "HBD": "N/A", "HBA": "N/A", "Violations": "N/A"}

def build_docx_report(metadata: Dict[str, Any], results: List[Dict[str, Any]], output_path: str):
    doc = docx.Document()
    title = doc.add_heading('Virtual Screening Pipeline Report (PyTorch DL Engine)', 0)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    doc.add_heading('1. Run Summary', level=1)
    p = doc.add_paragraph()
    p.add_run(f"Target Receptor: {metadata.get('receptor', 'N/A')}\n").bold = True
    p.add_run(f"Total Compounds Processed: {len(results)}\n")
    p.add_run(f"Execution Engine: PyTorch Deep Learning Model\n")

    doc.add_heading('2. Top Candidate Hits', level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rank'
    hdr_cells[1].text = 'Compound ID'
    hdr_cells[2].text = 'Affinity Score'
    hdr_cells[3].text = 'RO5 Violations'

    for i, res in enumerate(results[:10], 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = str(res.get('compound_id'))
        row_cells[2].text = str(res.get('affinity'))
        row_cells[3].text = str(res.get('admet', {}).get('Violations', 'N/A'))

    doc.save(output_path)

def rank_results(run_dir: str) -> List[Dict[str, Any]]:
    docking_out_dir = os.path.join(run_dir, "docking_out")
    summary_json = os.path.join(run_dir, "docking_summary.json")

    results = scan_and_rank_results(docking_out_dir, summary_json)
    for item in results:
        # Pass docked PDB/SDF path for ADMET
        mol_file = item.get("out_pdb") or item.get("out_sdf", "")
        item["admet"] = run_admet_analysis(mol_file)

    return results

def finalize_report(
    run_dir: str,
    corpus_dir: str,
    metadata: Dict[str, Any],
    results: List[Dict[str, Any]],
    top_n_for_rag: int = 5
) -> Dict[str, Any]:
    rag_handler = LiteratureRAGHandler(corpus_dir)

    for idx, item in enumerate(results):
        if idx < top_n_for_rag:
            item["rag_info"] = rag_handler.analyze_compound(item["compound_id"])
        else:
            item["rag_info"] = {
                "compound_id": item["compound_id"],
                "medical_uses": "Not analyzed (outside top hits).",
                "dosage_safety": "Unspecified.",
                "antiviral_findings": "Not analyzed (outside top hits)."
            }

    write_results_json(results, os.path.join(run_dir, "results.json"))
    write_results_csv(results, os.path.join(run_dir, "results.csv"))
    build_docx_report(metadata, results, os.path.join(run_dir, "screening_report.docx"))

    return {"status": "success", "total_processed": len(results), "results": results}
